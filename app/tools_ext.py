"""
CyberAgent 2.0 — herramientas de extensión (additivas).

Cada función es autocontenida y hace imports perezosos de sus dependencias,
de modo que una dependencia ausente NUNCA rompe la carga del módulo ni la app:
devuelve un error claro con instrucciones de instalación.

Incluye:
  - git_op          (GIT-001)    control de versiones
  - read_document   (INGEST-001) leer PDF/Excel/CSV/Word/código del usuario
  - browse_page     (BROWSER-001) navegador headless (Playwright)
  - schedule_task / list_scheduled / cancel_scheduled (SCHED-001)
"""
from __future__ import annotations

import os
import subprocess


# ════════════════════════════════════════════════════════════════════════════
# GIT-001 — control de versiones
# ════════════════════════════════════════════════════════════════════════════
_GIT_OPS = {
    "status": ["status", "--short", "--branch"],
    "log": ["log", "--oneline", "-15"],
    "diff": ["diff"],
    "branch": ["branch", "-a"],
    "pull": ["pull", "--ff-only"],
    "fetch": ["fetch", "--all"],
    "current_branch": ["rev-parse", "--abbrev-ref", "HEAD"],
    "remotes": ["remote", "-v"],
}


def git_op(operation: str, repo_path: str = ".", message: str = "",
           remote: str = "origin", branch: str = "", paths: list | None = None,
           url: str = "") -> dict:
    """Operaciones git seguras (whitelist). No usa shell; argumentos explícitos."""
    operation = (operation or "").strip().lower()
    repo = os.path.abspath(repo_path or ".")

    def _run(args: list, cwd: str | None = None) -> dict:
        try:
            r = subprocess.run(["git", *args], cwd=cwd or repo, capture_output=True,
                               text=True, timeout=120, encoding="utf-8", errors="replace")
            return {"ok": r.returncode == 0, "stdout": r.stdout[:8000],
                    "stderr": r.stderr[:2000], "returncode": r.returncode}
        except FileNotFoundError:
            return {"ok": False, "error": "git no está instalado o no está en PATH"}
        except subprocess.TimeoutExpired:
            return {"ok": False, "error": "git: timeout (120s)"}

    if operation == "clone":
        if not url:
            return {"ok": False, "error": "clone requiere 'url'"}
        dest = repo_path if repo_path not in (".", "") else os.path.basename(url).removesuffix(".git")
        return {"operation": "clone", **_run(["clone", url, dest], cwd=os.getcwd())}

    if not os.path.isdir(repo):
        return {"ok": False, "error": f"No existe el repo: {repo}"}

    if operation in _GIT_OPS:
        return {"operation": operation, "repo": repo, **_run(_GIT_OPS[operation])}

    if operation == "add":
        files = paths or ["-A"]
        return {"operation": "add", **_run(["add", *files])}

    if operation == "commit":
        if not message:
            return {"ok": False, "error": "commit requiere 'message'"}
        if paths:
            _run(["add", *paths])
        else:
            _run(["add", "-A"])
        return {"operation": "commit", **_run(["commit", "-m", message])}

    if operation == "push":
        args = ["push", remote]
        if branch:
            args.append(branch)
        return {"operation": "push", **_run(args)}

    if operation == "checkout":
        if not branch:
            return {"ok": False, "error": "checkout requiere 'branch'"}
        return {"operation": "checkout", **_run(["checkout", branch])}

    if operation == "create_branch":
        if not branch:
            return {"ok": False, "error": "create_branch requiere 'branch'"}
        return {"operation": "create_branch", **_run(["checkout", "-b", branch])}

    return {"ok": False, "error": f"operación git no soportada: {operation}",
            "soportadas": sorted(list(_GIT_OPS) + ["clone", "add", "commit", "push",
                                                   "checkout", "create_branch"])}


# ════════════════════════════════════════════════════════════════════════════
# INGEST-001 — leer documentos del usuario
# ════════════════════════════════════════════════════════════════════════════
_MAX_TEXT = 20000


def _clip(text: str, n: int = _MAX_TEXT) -> tuple[str, bool]:
    if len(text) <= n:
        return text, False
    return text[:n] + f"\n... [truncado: {len(text) - n} caracteres más]", True


def read_document(path: str, max_pages: int = 50) -> dict:
    """Lee y extrae texto de PDF, Word, Excel, CSV, JSON o código/texto plano."""
    if not os.path.isfile(path):
        return {"ok": False, "error": f"No existe el archivo: {path}"}
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    size = os.path.getsize(path)
    try:
        if ext == "pdf":
            return _read_pdf(path, max_pages, size)
        if ext in ("xlsx", "xlsm"):
            return _read_xlsx(path, size)
        if ext == "docx":
            return _read_docx(path, size)
        if ext == "csv":
            return _read_csv(path, size)
        # texto / código / json / md / etc.
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text, trunc = _clip(f.read())
        return {"ok": True, "type": ext or "text", "bytes": size,
                "text": text, "truncated": trunc}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "type": ext}


def _read_pdf(path: str, max_pages: int, size: int) -> dict:
    try:
        from pypdf import PdfReader
    except Exception:
        return {"ok": False, "error": "Falta 'pypdf'. Instala con: pip install pypdf",
                "type": "pdf"}
    reader = PdfReader(path)
    n = len(reader.pages)
    parts = []
    for i, page in enumerate(reader.pages[:max_pages]):
        parts.append(f"--- página {i + 1} ---\n" + (page.extract_text() or ""))
    text, trunc = _clip("\n\n".join(parts))
    return {"ok": True, "type": "pdf", "pages": n, "pages_read": min(n, max_pages),
            "bytes": size, "text": text, "truncated": trunc}


def _read_xlsx(path: str, size: int) -> dict:
    try:
        from openpyxl import load_workbook
    except Exception:
        return {"ok": False, "error": "Falta 'openpyxl'. Instala con: pip install openpyxl",
                "type": "xlsx"}
    wb = load_workbook(path, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"### Hoja: {ws.title}")
        for r, row in enumerate(ws.iter_rows(values_only=True)):
            if r >= 200:
                out.append("... [más filas omitidas]")
                break
            out.append(" | ".join("" if c is None else str(c) for c in row))
    text, trunc = _clip("\n".join(out))
    return {"ok": True, "type": "xlsx", "sheets": wb.sheetnames, "bytes": size,
            "text": text, "truncated": trunc}


def _read_docx(path: str, size: int) -> dict:
    try:
        import docx  # python-docx
    except Exception:
        return {"ok": False, "error": "Falta 'python-docx'. Instala con: pip install python-docx",
                "type": "docx"}
    d = docx.Document(path)
    text, trunc = _clip("\n".join(p.text for p in d.paragraphs))
    return {"ok": True, "type": "docx", "paragraphs": len(d.paragraphs),
            "bytes": size, "text": text, "truncated": trunc}


def _read_csv(path: str, size: int) -> dict:
    import csv
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
        for i, row in enumerate(csv.reader(f)):
            if i >= 500:
                rows.append("... [más filas omitidas]")
                break
            rows.append(" | ".join(row))
    text, trunc = _clip("\n".join(rows))
    return {"ok": True, "type": "csv", "rows_read": len(rows), "bytes": size,
            "text": text, "truncated": trunc}


# ════════════════════════════════════════════════════════════════════════════
# BROWSER-001 — navegador headless (Playwright)
# ════════════════════════════════════════════════════════════════════════════
def browse_page(url: str, action: str = "read", selector: str = "",
                text: str = "", wait_ms: int = 1500, screenshot: bool = False) -> dict:
    """
    Navegador headless real (ejecuta JS, SPAs, login, formularios).
    action: read | click | fill | screenshot
    Requiere: pip install playwright && playwright install chromium
    """
    try:
        from playwright.sync_api import sync_playwright
    except Exception:
        return {"ok": False,
                "error": "Falta Playwright. Instala con:\n"
                         "  pip install playwright\n  playwright install chromium",
                "fallback_hint": "Mientras tanto puedes usar web_fetch (sin JS)."}
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, wait_until="domcontentloaded", timeout=30000)
            page.wait_for_timeout(min(max(int(wait_ms), 0), 10000))

            result: dict = {"ok": True, "url": page.url, "title": page.title(), "action": action}

            if action == "fill" and selector:
                page.fill(selector, text)
                result["filled"] = selector
            elif action == "click" and selector:
                page.click(selector, timeout=8000)
                page.wait_for_timeout(800)
                result["clicked"] = selector
                result["url"] = page.url

            if action == "read" or action != "screenshot":
                body = page.inner_text("body")
                if len(body) > _MAX_TEXT:
                    body = body[:_MAX_TEXT] + "\n... [truncado]"
                result["text"] = body
                links = page.eval_on_selector_all(
                    "a[href]", "els => els.slice(0,40).map(a => ({t:a.innerText.trim().slice(0,60), h:a.href}))")
                result["links"] = [l for l in links if l.get("t")][:30]

            if screenshot or action == "screenshot":
                from app.documents import SERVED_DIR, public_url_for
                os.makedirs(SERVED_DIR, exist_ok=True)
                import time as _t
                fname = f"page_{int(_t.time())}.png"
                fpath = os.path.join(SERVED_DIR, fname)
                page.screenshot(path=fpath, full_page=True)
                result["screenshot_url"] = public_url_for(fpath)

            browser.close()
            return result
    except Exception as exc:
        return {"ok": False, "error": str(exc), "url": url}


# ════════════════════════════════════════════════════════════════════════════
# MSG-001 — mensajería saliente (email + Telegram)
# ════════════════════════════════════════════════════════════════════════════
def send_message(channel: str, text: str, to: str = "", subject: str = "") -> dict:
    """
    Envía un mensaje por email (SMTP) o Telegram para entregar resultados/avisos.
    Config por variables de entorno:
      Email:    CYBERAGENT_SMTP_HOST, _PORT(=587), _USER, _PASS, _FROM
      Telegram: CYBERAGENT_TELEGRAM_TOKEN, CYBERAGENT_TELEGRAM_CHAT_ID
    """
    channel = (channel or "").strip().lower()
    if not text:
        return {"ok": False, "error": "text vacío"}
    if channel == "email":
        return _send_email(to, subject or "CyberAgent", text)
    if channel == "telegram":
        return _send_telegram(text, to)
    return {"ok": False, "error": "channel debe ser 'email' o 'telegram'"}


def _send_email(to: str, subject: str, body: str) -> dict:
    host = os.getenv("CYBERAGENT_SMTP_HOST", "")
    user = os.getenv("CYBERAGENT_SMTP_USER", "")
    pw = os.getenv("CYBERAGENT_SMTP_PASS", "")
    sender = os.getenv("CYBERAGENT_SMTP_FROM", user)
    port = int(os.getenv("CYBERAGENT_SMTP_PORT", "587"))
    to = to or os.getenv("CYBERAGENT_SMTP_TO", sender)
    if not (host and user and pw):
        return {"ok": False, "error": "Email no configurado. Define CYBERAGENT_SMTP_HOST/_USER/_PASS"
                                      " (y opcional _PORT, _FROM, _TO) en el entorno."}
    if not to:
        return {"ok": False, "error": "falta destinatario 'to' (o CYBERAGENT_SMTP_TO)"}
    try:
        import smtplib
        from email.mime.text import MIMEText
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = sender
        msg["To"] = to
        with smtplib.SMTP(host, port, timeout=30) as s:
            s.starttls()
            s.login(user, pw)
            s.sendmail(sender, [a.strip() for a in to.split(",")], msg.as_string())
        return {"ok": True, "channel": "email", "to": to, "subject": subject}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "channel": "email"}


def _send_telegram(text: str, chat_id: str = "") -> dict:
    token = os.getenv("CYBERAGENT_TELEGRAM_TOKEN", "")
    chat_id = chat_id or os.getenv("CYBERAGENT_TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        return {"ok": False, "error": "Telegram no configurado. Define CYBERAGENT_TELEGRAM_TOKEN y "
                                      "CYBERAGENT_TELEGRAM_CHAT_ID en el entorno."}
    try:
        import httpx
        r = httpx.post(f"https://api.telegram.org/bot{token}/sendMessage",
                       json={"chat_id": chat_id, "text": text[:4000]}, timeout=20)
        if r.status_code != 200:
            return {"ok": False, "error": f"Telegram HTTP {r.status_code}: {r.text[:200]}"}
        return {"ok": True, "channel": "telegram", "chat_id": chat_id}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "channel": "telegram"}


# ════════════════════════════════════════════════════════════════════════════
# VERIFY-001 — herramientas de verificación de red (para COMPROBAR, no asumir)
# ════════════════════════════════════════════════════════════════════════════
def http_check(url: str, method: str = "GET", expect_status=None, timeout: int = 15) -> dict:
    """Comprueba que una URL responde: código HTTP, tiempo y trozo del cuerpo."""
    import time
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible (pip install httpx)"}
    try:
        t0 = time.time()
        r = httpx.request((method or "GET").upper(), url, timeout=timeout, follow_redirects=True)
        dt = round(time.time() - t0, 3)
        ok = True if expect_status is None else (int(r.status_code) == int(expect_status))
        return {"ok": ok, "status": r.status_code, "elapsed_s": dt,
                "final_url": str(r.url), "expected": expect_status,
                "body_preview": (r.text or "")[:300]}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}", "url": url}


def dns_resolve(host: str) -> dict:
    """Resuelve un dominio a IPs (A/AAAA). NXDOMAIN = el dominio no existe."""
    import socket
    h = (host or "").strip().replace("https://", "").replace("http://", "").split("/")[0]
    if not h:
        return {"ok": False, "error": "host vacío"}
    try:
        infos = socket.getaddrinfo(h, None)
        ips = sorted({i[4][0] for i in infos})
        return {"ok": True, "host": h, "ips": ips}
    except Exception as e:
        return {"ok": False, "host": h, "error": f"{type(e).__name__}: {e}",
                "hint": "Si es 'no such host'/NXDOMAIN, el dominio no existe o el DNS no ha propagado."}


def port_check(host: str, port: int, timeout: int = 5) -> dict:
    """Comprueba si un puerto TCP está abierto (servicio escuchando)."""
    import socket
    h = (host or "").strip().replace("https://", "").replace("http://", "").split("/")[0]
    try:
        with socket.create_connection((h, int(port)), timeout=timeout):
            return {"ok": True, "host": h, "port": int(port), "open": True}
    except Exception as e:
        return {"ok": True, "host": h, "port": int(port), "open": False,
                "error": f"{type(e).__name__}: {e}"}


# ════════════════════════════════════════════════════════════════════════════
# MISTRAL-API-001 — endpoints nativos de Mistral que faltaban (OCR, visión,
# audio, embeddings, moderación, código). Cada uno autocontenido y defensivo.
# ════════════════════════════════════════════════════════════════════════════
def _m_key() -> str:
    from app.brain import mistral_api_key
    return mistral_api_key()

def _m_base() -> str:
    from app.brain import MISTRAL_BASE_URL
    return MISTRAL_BASE_URL

def _m_headers() -> dict:
    return {"Authorization": f"Bearer {_m_key()}", "Content-Type": "application/json"}

def _m_log(model, usage, ctx):
    try:
        from app import mistral_usage
        usage = usage or {}
        mistral_usage.log_usage(model, usage.get("prompt_tokens", 0),
                                usage.get("completion_tokens", usage.get("total_tokens", 0)), ctx)
    except Exception:
        pass

def _as_data_url(path_or_url: str, kind: str = "image"):
    """Devuelve una URL http(s) tal cual, o un data: base64 si es un archivo local."""
    s = (path_or_url or "").strip()
    if s.startswith("http://") or s.startswith("https://"):
        return s
    if os.path.isfile(s):
        import base64, mimetypes
        mime = mimetypes.guess_type(s)[0] or ("image/png" if kind == "image" else "application/pdf")
        with open(s, "rb") as f:
            return "data:" + mime + ";base64," + base64.b64encode(f.read()).decode()
    return None


def mistral_ocr(source: str) -> dict:
    """OCR real de Mistral sobre un documento/imagen (URL o archivo local)."""
    if not _m_key():
        return {"ok": False, "error": "MISTRAL_API_KEY no configurada"}
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    url = _as_data_url(source, "doc")
    if not url:
        return {"ok": False, "error": "no encuentro la fuente: " + str(source)}
    _imgext = (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp")
    _base = url.lower().split("?")[0]
    is_image = url.startswith("data:image") or any(_base.endswith(e) for e in _imgext)
    doc = ({"type": "image_url", "image_url": url} if is_image
           else {"type": "document_url", "document_url": url})
    try:
        r = httpx.post(_m_base() + "/ocr", headers=_m_headers(), timeout=120,
                       json={"model": "mistral-ocr-latest", "document": doc})
        if r.status_code >= 400:
            return {"ok": False, "error": "HTTP " + str(r.status_code) + ": " + r.text[:300]}
        data = r.json()
        pages = data.get("pages", [])
        text = "\n\n".join(p.get("markdown", "") for p in pages)
        _m_log("mistral-ocr-latest", data.get("usage_info") or data.get("usage"), "ocr")
        return {"ok": True, "pages": len(pages), "text": text[:12000]}
    except Exception as e:
        return {"ok": False, "error": type(e).__name__ + ": " + str(e)}


def mistral_vision(image: str, question: str = "Describe la imagen en detalle.") -> dict:
    """Entiende una imagen con Pixtral (URL o archivo local) y responde una pregunta."""
    if not _m_key():
        return {"ok": False, "error": "MISTRAL_API_KEY no configurada"}
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    url = _as_data_url(image, "image")
    if not url:
        return {"ok": False, "error": "no encuentro la imagen: " + str(image)}
    try:
        r = httpx.post(_m_base() + "/chat/completions", headers=_m_headers(), timeout=90,
                       json={"model": "pixtral-large-latest", "messages": [
                           {"role": "user", "content": [
                               {"type": "text", "text": question},
                               {"type": "image_url", "image_url": {"url": url}}]}]})
        if r.status_code >= 400:
            return {"ok": False, "error": "HTTP " + str(r.status_code) + ": " + r.text[:300]}
        data = r.json()
        _m_log("pixtral-large-latest", data.get("usage"), "vision")
        return {"ok": True, "answer": data["choices"][0]["message"]["content"]}
    except Exception as e:
        return {"ok": False, "error": type(e).__name__ + ": " + str(e)}


def mistral_transcribe(audio_path: str, language: str = "") -> dict:
    """Transcribe audio con Voxtral (archivo local)."""
    if not _m_key():
        return {"ok": False, "error": "MISTRAL_API_KEY no configurada"}
    if not os.path.isfile(audio_path):
        return {"ok": False, "error": "archivo no encontrado: " + str(audio_path)}
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    try:
        data = {"model": "voxtral-mini-latest"}
        if language:
            data["language"] = language
        with open(audio_path, "rb") as f:
            r = httpx.post(_m_base() + "/audio/transcriptions",
                           headers={"Authorization": "Bearer " + _m_key()},
                           data=data, files={"file": (os.path.basename(audio_path), f)}, timeout=180)
        if r.status_code >= 400:
            return {"ok": False, "error": "HTTP " + str(r.status_code) + ": " + r.text[:300]}
        return {"ok": True, "text": r.json().get("text", "")}
    except Exception as e:
        return {"ok": False, "error": type(e).__name__ + ": " + str(e)}


def mistral_embed(texts) -> dict:
    """Embeddings con mistral-embed. Con 2 textos, devuelve su similitud coseno."""
    if not _m_key():
        return {"ok": False, "error": "MISTRAL_API_KEY no configurada"}
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    if isinstance(texts, str):
        texts = [texts]
    try:
        r = httpx.post(_m_base() + "/embeddings", headers=_m_headers(), timeout=60,
                       json={"model": "mistral-embed", "input": texts})
        if r.status_code >= 400:
            return {"ok": False, "error": "HTTP " + str(r.status_code) + ": " + r.text[:300]}
        data = r.json()
        vecs = [d["embedding"] for d in data.get("data", [])]
        out = {"ok": True, "count": len(vecs), "dim": len(vecs[0]) if vecs else 0}
        if len(vecs) == 2:
            import math
            a, b = vecs
            dot = sum(x * y for x, y in zip(a, b))
            na = math.sqrt(sum(x * x for x in a)); nb = math.sqrt(sum(y * y for y in b))
            out["cosine_similarity"] = round(dot / (na * nb), 4) if na and nb else None
        return out
    except Exception as e:
        return {"ok": False, "error": type(e).__name__ + ": " + str(e)}


def mistral_moderate(text: str) -> dict:
    """Clasifica contenido con la API de moderación de Mistral."""
    if not _m_key():
        return {"ok": False, "error": "MISTRAL_API_KEY no configurada"}
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    try:
        r = httpx.post(_m_base() + "/moderations", headers=_m_headers(), timeout=30,
                       json={"model": "mistral-moderation-latest", "input": [text]})
        if r.status_code >= 400:
            return {"ok": False, "error": "HTTP " + str(r.status_code) + ": " + r.text[:300]}
        res = r.json().get("results", [{}])[0]
        cats = {k: v for k, v in (res.get("categories") or {}).items() if v}
        return {"ok": True, "flagged_categories": list(cats.keys()), "raw": res}
    except Exception as e:
        return {"ok": False, "error": type(e).__name__ + ": " + str(e)}


def mistral_code_complete(prompt: str, suffix: str = "") -> dict:
    """Completa código con Codestral (FIM: prompt + suffix opcional)."""
    if not _m_key():
        return {"ok": False, "error": "MISTRAL_API_KEY no configurada"}
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    try:
        body = {"model": "codestral-latest", "prompt": prompt, "max_tokens": 512}
        if suffix:
            body["suffix"] = suffix
        r = httpx.post(_m_base() + "/fim/completions", headers=_m_headers(), timeout=60, json=body)
        if r.status_code >= 400:
            return {"ok": False, "error": "HTTP " + str(r.status_code) + ": " + r.text[:300]}
        data = r.json()
        _m_log("codestral-latest", data.get("usage"), "fim")
        return {"ok": True, "completion": data["choices"][0]["message"]["content"]}
    except Exception as e:
        return {"ok": False, "error": type(e).__name__ + ": " + str(e)}


# ════════════════════════════════════════════════════════════════════════════
# EDIT-001 — edición quirúrgica de archivos (find/replace). El cambio clave para
# que el agente programe bien: edita trozos en vez de reescribir archivos enteros.
# ════════════════════════════════════════════════════════════════════════════
def _make_unified_diff(before: str, after: str, path: str, max_lines: int = 120) -> str:
    """Diff unificado estilo git para que la UI lo pinte como Claude Code."""
    import difflib
    diff = difflib.unified_diff(
        before.splitlines(), after.splitlines(),
        fromfile="a/" + os.path.basename(path), tofile="b/" + os.path.basename(path),
        lineterm="", n=3,
    )
    lines = list(diff)
    if len(lines) > max_lines:
        lines = lines[:max_lines] + [f"... (+{len(lines) - max_lines} líneas de diff omitidas)"]
    return "\n".join(lines)


def _find_occurrences(content: str, needle: str) -> int:
    return content.count(needle)


def _read_with_eol(path: str):
    """Lee el archivo y devuelve (content_lf, eol). content_lf usa siempre \\n."""
    with open(path, "r", encoding="utf-8", newline="") as f:
        raw = f.read()
    crlf = raw.count("\r\n")
    lf_only = raw.count("\n") - crlf
    eol = "\r\n" if crlf > lf_only else "\n"
    content = raw.replace("\r\n", "\n").replace("\r", "\n")
    return content, eol


def _write_with_eol(path: str, content_lf: str, eol: str):
    out = content_lf.replace("\n", eol) if eol == "\r\n" else content_lf
    with open(path, "w", encoding="utf-8", newline="") as f:
        f.write(out)
    return len(out)


def _apply_replacement(content: str, old_string: str, new_string: str, replace_all: bool):
    """Aplica un reemplazo sobre content (LF). Tolera CRLF/LF y espacios finales.
    Devuelve (new_content, count, mode). Lanza ValueError con mensaje claro si falla."""
    if old_string == new_string:
        raise ValueError("old_string y new_string son iguales")
    old_n = old_string.replace("\r\n", "\n").replace("\r", "\n")
    new_n = new_string.replace("\r\n", "\n").replace("\r", "\n")

    mode = "exacto"
    count = content.count(old_n)
    if count == 0:
        # Fallback: tolerar espacios/tabs finales en cada línea del bloque.
        def _strip_trailing(s: str) -> str:
            return "\n".join(ln.rstrip() for ln in s.split("\n"))
        content_st = _strip_trailing(content)
        old_st = _strip_trailing(old_n)
        if old_st and content_st.count(old_st) == 1:
            idx = content_st.index(old_st)
            line_start = content_st.count("\n", 0, idx)
            n_lines = old_st.count("\n") + 1
            old_n = "\n".join(content.split("\n")[line_start:line_start + n_lines])
            count = 1
            mode = "espacios-finales-tolerados"
    if count == 0:
        # Red de seguridad: si el modelo pegó los números de línea de read_file
        # (formato "  123\t código"), quítalos y reintenta.
        import re as _re
        stripped = "\n".join(_re.sub(r"^\s*\d+\t", "", ln) for ln in old_n.split("\n"))
        if stripped != old_n and content.count(stripped) == 1:
            old_n = stripped
            count = 1
            mode = "num-linea-eliminado"
    if count == 0:
        raise ValueError("no se encontró old_string (copia el texto EXACTO con read_file; "
                         "CRLF/LF ya se toleran): " + repr(old_string[:60]))
    if count > 1 and not replace_all:
        raise ValueError(f"old_string aparece {count} veces; añade contexto único o replace_all=true")
    new_content = (content.replace(old_n, new_n) if replace_all
                   else content.replace(old_n, new_n, 1))
    return new_content, (count if replace_all else 1), mode


def _py_syntax_result(path: str, content_lf: str, result: dict) -> dict:
    if path.lower().endswith(".py"):
        try:
            import ast
            ast.parse(content_lf)
            result["syntax_ok"] = True
        except SyntaxError as se:
            result["syntax_ok"] = False
            result["syntax_error"] = f"línea {se.lineno}: {se.msg}"
            result["warning"] = ("⚠️ El archivo quedó editado pero ya NO compila. "
                                 "Revisa el diff y corrige.")
    return result


def edit_file(path: str, old_string: str, new_string: str, replace_all: bool = False) -> dict:
    """Reemplaza old_string por new_string en el archivo. old_string debe ser
    ÚNICO (añade contexto si no), salvo replace_all=True. No reescribe el archivo
    entero: solo cambia el trozo indicado, así es fiable para código.

    Robusto en Windows: tolera CRLF/LF y espacios finales, preserva el fin de línea
    original y devuelve un diff unificado del cambio."""
    if not os.path.isfile(path):
        return {"ok": False, "error": "no existe el archivo: " + str(path)}
    try:
        content, eol = _read_with_eol(path)
    except Exception as e:
        return {"ok": False, "error": "lectura: " + str(e)}
    try:
        new_content, count, mode = _apply_replacement(content, old_string, new_string, replace_all)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    diff = _make_unified_diff(content, new_content, path)
    try:
        size = _write_with_eol(path, new_content, eol)
    except Exception as e:
        return {"ok": False, "error": "escritura: " + str(e)}
    result = {"ok": True, "path": path, "replacements": count,
              "new_size": size, "match_mode": mode, "diff": diff}
    return _py_syntax_result(path, new_content, result)


def lint_code(path: str, security: bool = False) -> dict:
    """Analiza un archivo con linters de CPU (gratis, NO usan GPU → no compiten con
    la inferencia del modelo) y devuelve los problemas reales para que el modelo se
    AUTOCORRIJA. Úsalo tras editar código.
      • Python: ruff (estilo + bugs) y, si security=True, bandit (seguridad).
      • JS/TS: node --check (sintaxis).
    Devuelve {ok, issues:[{line, code, message, severity}], clean}."""
    import json as _json
    import subprocess as _sp
    import sys as _sys
    if not os.path.isfile(path):
        return {"ok": False, "error": "no existe el archivo: " + str(path)}
    ext = os.path.splitext(path)[1].lower()
    issues = []
    tools_used = []

    def _run(cmd, timeout=60):
        try:
            return _sp.run(cmd, capture_output=True, text=True, timeout=timeout)
        except Exception as e:
            return type("R", (), {"returncode": -1, "stdout": "", "stderr": str(e)})()

    if ext == ".py":
        r = _run([_sys.executable, "-m", "ruff", "check", "--output-format", "json", path])
        tools_used.append("ruff")
        try:
            for d in _json.loads(r.stdout or "[]"):
                loc = d.get("location") or {}
                issues.append({"line": loc.get("row"), "code": d.get("code"),
                               "message": d.get("message"), "severity": "warning"})
        except Exception:
            if r.stderr.strip():
                issues.append({"line": None, "code": "ruff", "message": r.stderr[:300],
                               "severity": "error"})
        if security:
            rb = _run([_sys.executable, "-m", "bandit", "-f", "json", "-q", path])
            tools_used.append("bandit")
            try:
                for d in (_json.loads(rb.stdout or "{}").get("results") or []):
                    issues.append({"line": d.get("line_number"),
                                   "code": d.get("test_id"),
                                   "message": d.get("issue_text"),
                                   "severity": (d.get("issue_severity") or "low").lower()})
            except Exception:
                pass
    elif ext in (".js", ".mjs", ".cjs", ".ts"):
        r = _run(["node", "--check", path])
        tools_used.append("node --check")
        if r.returncode != 0 and r.stderr.strip():
            issues.append({"line": None, "code": "syntax",
                           "message": r.stderr.strip()[:300], "severity": "error"})
    else:
        return {"ok": False, "error": f"sin linter para {ext or 'este tipo'} "
                "(soportado: .py, .js/.ts)"}

    return {"ok": True, "path": path, "tools": tools_used,
            "issues": issues[:50], "issue_count": len(issues),
            "clean": len(issues) == 0}


# ════════════════════════════════════════════════════════════════════════════
# OLEADA DE HERRAMIENTAS (jun-2026) — código, seguridad, datos. Todo CPU/red/API,
# nada usa GPU (no compite con la inferencia). Los wrappers de binarios externos
# degradan con un mensaje claro si la herramienta no está instalada.
# ════════════════════════════════════════════════════════════════════════════

def _which(name: str):
    import shutil
    return shutil.which(name)


def _run_cmd(cmd, timeout=120, cwd=None):
    import subprocess
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, cwd=cwd)
        return r.returncode, r.stdout or "", r.stderr or ""
    except FileNotFoundError:
        return -1, "", "comando no encontrado: " + (cmd[0] if isinstance(cmd, list) else str(cmd))
    except subprocess.TimeoutExpired:
        return -2, "", f"timeout tras {timeout}s"
    except Exception as e:
        return -1, "", f"{type(e).__name__}: {e}"


# ── CÓDIGO ──────────────────────────────────────────────────────────────────
def run_tests(path: str = ".", framework: str = "auto", k: str = "") -> dict:
    """Ejecuta la batería de tests y devuelve pass/fail parseado para que el modelo
    cierre el bucle editar→testear→corregir. framework: auto|pytest|jest|go."""
    import sys
    if framework == "auto":
        if os.path.exists(os.path.join(path, "package.json")):
            framework = "jest"
        elif any(f.endswith(".go") for f in os.listdir(path)) if os.path.isdir(path) else False:
            framework = "go"
        else:
            framework = "pytest"
    if framework == "pytest":
        cmd = [sys.executable, "-m", "pytest", "-q", path]
        if k:
            cmd += ["-k", k]
    elif framework == "jest":
        cmd = ["npx", "jest", "--silent"]
        if k:
            cmd += ["-t", k]
    elif framework == "go":
        cmd = ["go", "test", "./..."]
    else:
        return {"ok": False, "error": "framework no soportado: " + framework}
    rc, out, err = _run_cmd(cmd, timeout=300, cwd=path if os.path.isdir(path) else None)
    tail = (out + "\n" + err).strip()
    import re as _re
    m = _re.search(r"(\d+) passed", tail)
    f = _re.search(r"(\d+) failed", tail)
    return {"ok": rc == 0, "framework": framework, "exit_code": rc,
            "passed": int(m.group(1)) if m else None,
            "failed": int(f.group(1)) if f else (0 if rc == 0 else None),
            "output": tail[-3000:]}


def apply_patch(patch: str) -> dict:
    """Aplica un diff unificado (git) que puede tocar VARIOS archivos, de forma
    atómica (git apply: si no encaja, no cambia nada). Ideal para refactors grandes."""
    import tempfile
    if not _which("git"):
        return {"ok": False, "error": "git no está instalado (necesario para apply_patch)"}
    if not patch.strip():
        return {"ok": False, "error": "patch vacío"}
    if not patch.endswith("\n"):
        patch += "\n"
    fd, tmp = tempfile.mkstemp(suffix=".patch", text=True)
    try:
        with os.fdopen(fd, "w", encoding="utf-8", newline="\n") as f:
            f.write(patch)
        rc, out, err = _run_cmd(["git", "apply", "--check", "--recount", tmp], timeout=30)
        if rc != 0:
            return {"ok": False, "error": "el patch no encaja: " + (err or out)[:400]}
        rc, out, err = _run_cmd(["git", "apply", "--recount", tmp], timeout=30)
        if rc != 0:
            return {"ok": False, "error": "fallo al aplicar: " + (err or out)[:400]}
        files = [ln[6:] for ln in patch.splitlines() if ln.startswith("+++ b/")]
        return {"ok": True, "files_changed": files, "count": len(files)}
    finally:
        try:
            os.remove(tmp)
        except Exception:
            pass


def code_symbols(path: str, find: str = "") -> dict:
    """Lista símbolos (clases/funciones) de un archivo Python con su línea — el modelo
    se ubica sin leer el archivo entero. Si `find`, devuelve también dónde se USA ese
    nombre (línea + texto), como find_references."""
    if not os.path.isfile(path):
        return {"ok": False, "error": "no existe: " + str(path)}
    try:
        src = open(path, "r", encoding="utf-8", errors="replace").read()
    except Exception as e:
        return {"ok": False, "error": str(e)}
    symbols = []
    if path.endswith(".py"):
        import ast
        try:
            tree = ast.parse(src)
            for node in ast.walk(tree):
                if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
                    kind = "class" if isinstance(node, ast.ClassDef) else "func"
                    symbols.append({"name": node.name, "kind": kind, "line": node.lineno})
        except SyntaxError as e:
            return {"ok": False, "error": f"sintaxis L{e.lineno}: {e.msg}"}
    else:
        import re as _re
        for i, ln in enumerate(src.splitlines(), 1):
            m = _re.match(r"\s*(?:export\s+)?(?:async\s+)?function\s+(\w+)|\s*(?:export\s+)?class\s+(\w+)", ln)
            if m:
                symbols.append({"name": m.group(1) or m.group(2),
                                "kind": "class" if m.group(2) else "func", "line": i})
    refs = []
    if find:
        for i, ln in enumerate(src.splitlines(), 1):
            if find in ln:
                refs.append({"line": i, "text": ln.strip()[:120]})
    out = {"ok": True, "path": path, "symbols": symbols[:200], "count": len(symbols)}
    if find:
        out["references"] = refs[:100]
    return out


# ── SEGURIDAD / INTEL (gratis, red) ─────────────────────────────────────────
def cve_lookup(query: str) -> dict:
    """Consulta la API pública de NVD (gratis, sin key). query = 'CVE-2024-XXXX'
    para uno concreto, o palabras clave para buscar. Devuelve descripción y CVSS."""
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    q = query.strip()
    base = "https://services.nvd.nist.gov/rest/json/cves/2.0"
    params = {"cveId": q.upper()} if q.upper().startswith("CVE-") else {"keywordSearch": q, "resultsPerPage": 5}
    try:
        r = httpx.get(base, params=params, timeout=25)
        if r.status_code != 200:
            return {"ok": False, "error": f"NVD HTTP {r.status_code}"}
        data = r.json()
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}
    out = []
    for item in (data.get("vulnerabilities") or [])[:5]:
        c = item.get("cve", {})
        desc = next((d["value"] for d in c.get("descriptions", []) if d.get("lang") == "en"), "")
        metrics = c.get("metrics", {})
        score, sev = None, None
        for key in ("cvssMetricV31", "cvssMetricV30", "cvssMetricV2"):
            if metrics.get(key):
                cd = metrics[key][0].get("cvssData", {})
                score = cd.get("baseScore")
                sev = cd.get("baseSeverity") or metrics[key][0].get("baseSeverity")
                break
        out.append({"id": c.get("id"), "cvss": score, "severity": sev,
                    "published": c.get("published"), "description": desc[:400]})
    return {"ok": True, "count": len(out), "results": out}


def threat_intel(indicator: str, kind: str = "auto") -> dict:
    """Reputación de un IOC (IP / hash / dominio / URL) vía VirusTotal o AbuseIPDB.
    Requiere key gratis: VIRUSTOTAL_API_KEY y/o ABUSEIPDB_API_KEY en el entorno."""
    import re as _re
    ind = indicator.strip()
    if kind == "auto":
        if _re.fullmatch(r"\d{1,3}(\.\d{1,3}){3}", ind):
            kind = "ip"
        elif _re.fullmatch(r"[A-Fa-f0-9]{32,64}", ind):
            kind = "hash"
        elif ind.startswith("http"):
            kind = "url"
        else:
            kind = "domain"
    try:
        import httpx
    except Exception:
        return {"ok": False, "error": "httpx no disponible"}
    vt = os.getenv("VIRUSTOTAL_API_KEY") or os.getenv("VT_API_KEY")
    abuse = os.getenv("ABUSEIPDB_API_KEY")
    try:
        if kind == "ip" and abuse:
            r = httpx.get("https://api.abuseipdb.com/api/v2/check",
                          headers={"Key": abuse, "Accept": "application/json"},
                          params={"ipAddress": ind, "maxAgeInDays": 90}, timeout=20)
            d = r.json().get("data", {})
            return {"ok": True, "source": "abuseipdb", "indicator": ind,
                    "abuse_score": d.get("abuseConfidenceScore"), "country": d.get("countryCode"),
                    "total_reports": d.get("totalReports"), "isp": d.get("isp")}
        if vt:
            ep = {"ip": "ip_addresses", "hash": "files", "domain": "domains", "url": "urls"}[kind]
            target = ind
            if kind == "url":
                import base64
                target = base64.urlsafe_b64encode(ind.encode()).decode().strip("=")
            r = httpx.get(f"https://www.virustotal.com/api/v3/{ep}/{target}",
                          headers={"x-apikey": vt}, timeout=20)
            if r.status_code == 404:
                return {"ok": True, "source": "virustotal", "indicator": ind, "found": False}
            stats = r.json().get("data", {}).get("attributes", {}).get("last_analysis_stats", {})
            return {"ok": True, "source": "virustotal", "indicator": ind, "kind": kind,
                    "malicious": stats.get("malicious"), "suspicious": stats.get("suspicious"),
                    "harmless": stats.get("harmless"), "stats": stats}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}
    return {"ok": False, "error": "falta API key (VIRUSTOTAL_API_KEY o ABUSEIPDB_API_KEY). "
            "Ambas tienen plan gratuito."}


def yara_scan(path: str, rules: str) -> dict:
    """Escanea un archivo/carpeta con reglas YARA (CPU). `rules` = reglas YARA (texto)
    o ruta a un .yar. Detección de malware/IOCs. Requiere: pip install yara-python."""
    try:
        import yara
    except Exception:
        return {"ok": False, "error": "yara-python no instalado: pip install yara-python"}
    try:
        compiled = yara.compile(filepath=rules) if os.path.isfile(rules) else yara.compile(source=rules)
    except Exception as e:
        return {"ok": False, "error": "reglas inválidas: " + str(e)[:200]}
    targets = []
    if os.path.isdir(path):
        for root, _, files in os.walk(path):
            for fn in files:
                targets.append(os.path.join(root, fn))
    elif os.path.isfile(path):
        targets = [path]
    else:
        return {"ok": False, "error": "no existe: " + str(path)}
    matches = []
    for t in targets[:500]:
        try:
            for m in compiled.match(t, timeout=20):
                matches.append({"file": t, "rule": m.rule,
                                "tags": list(m.tags), "strings": len(m.strings)})
        except Exception:
            continue
    return {"ok": True, "scanned": len(targets), "matches": matches[:100],
            "match_count": len(matches), "clean": len(matches) == 0}


# ── RECON OFENSIVO (wrappers de binarios reales) ────────────────────────────
def nmap_scan(target: str, options: str = "-sV -T4 --top-ports 100") -> dict:
    """Escaneo con el nmap REAL (más potente que port_scan). options: flags de nmap.
    Requiere nmap instalado. Uso autorizado de pentesting."""
    if not _which("nmap"):
        return {"ok": False, "error": "nmap no instalado (https://nmap.org/download)"}
    import shlex
    cmd = ["nmap"] + shlex.split(options) + [target]
    rc, out, err = _run_cmd(cmd, timeout=300)
    return {"ok": rc == 0, "target": target, "output": (out or err)[-6000:]}


def web_audit(target: str, tool: str = "nikto") -> dict:
    """Auditoría web ofensiva con la herramienta real. tool: nikto | sqlmap | ffuf.
    Para ffuf usa el placeholder FUZZ en el target. Uso autorizado."""
    if tool == "nikto":
        if not _which("nikto"):
            return {"ok": False, "error": "nikto no instalado"}
        cmd = ["nikto", "-h", target, "-maxtime", "120s"]
    elif tool == "sqlmap":
        if not _which("sqlmap"):
            return {"ok": False, "error": "sqlmap no instalado (pip install sqlmap)"}
        cmd = ["sqlmap", "-u", target, "--batch", "--smart", "--level", "1"]
    elif tool == "ffuf":
        if not _which("ffuf"):
            return {"ok": False, "error": "ffuf no instalado"}
        cmd = ["ffuf", "-u", target, "-w", os.getenv("FFUF_WORDLIST", "/usr/share/wordlists/common.txt"), "-mc", "200,301,302,403"]
    else:
        return {"ok": False, "error": "tool no soportada: " + tool}
    rc, out, err = _run_cmd(cmd, timeout=240)
    return {"ok": rc in (0, 1), "tool": tool, "target": target, "output": (out or err)[-6000:]}


def hash_crack(target: str, wordlist: str = "", mode: str = "", tool: str = "john") -> dict:
    """Crackea hashes (pentesting autorizado). target = archivo con hashes. tool:
    john (CPU, NO usa GPU → no compite con el modelo) | hashcat (GPU, COMPETIRÍA con
    la inferencia). mode = formato john / -m de hashcat. Requiere john o hashcat."""
    if not os.path.isfile(target):
        return {"ok": False, "error": "el target debe ser un archivo con hashes"}
    wl = wordlist or os.getenv("CRACK_WORDLIST", "")
    if tool == "john":
        if not _which("john"):
            return {"ok": False, "error": "John the Ripper no instalado"}
        cmd = ["john"]
        if wl:
            cmd.append("--wordlist=" + wl)
        if mode:
            cmd.append("--format=" + mode)
        cmd.append(target)
        rc, out, err = _run_cmd(cmd, timeout=300)
        _, shown, _ = _run_cmd(["john", "--show", target], timeout=30)
        return {"ok": True, "tool": "john", "cracked": shown[-3000:], "log": (out or err)[-1500:],
                "note": "john usa CPU → no compite con la GPU del modelo"}
    elif tool == "hashcat":
        if not _which("hashcat"):
            return {"ok": False, "error": "hashcat no instalado"}
        if not (wl and mode):
            return {"ok": False, "error": "hashcat requiere wordlist y mode (-m)"}
        cmd = ["hashcat", "-m", mode, "-a", "0", target, wl, "--quiet"]
        rc, out, err = _run_cmd(cmd, timeout=300)
        return {"ok": rc in (0, 1), "tool": "hashcat", "output": (out or err)[-4000:],
                "warning": "⚠️ hashcat usa la GPU → ralentiza la inferencia del modelo mientras corre"}
    return {"ok": False, "error": "tool no soportada: " + tool}


# ── DATOS ───────────────────────────────────────────────────────────────────
def sql_query(db_path: str, query: str, allow_write: bool = False) -> dict:
    """Ejecuta SQL sobre una base SQLite local. Por seguridad solo permite SELECT/
    PRAGMA salvo allow_write=True. CPU, gratis."""
    import sqlite3
    if not os.path.isfile(db_path):
        return {"ok": False, "error": "no existe la base: " + str(db_path)}
    q = query.strip()
    if not allow_write and not q.lower().lstrip("(").startswith(("select", "pragma", "with", "explain")):
        return {"ok": False, "error": "solo lectura (SELECT/PRAGMA). Usa allow_write=true para modificar."}
    try:
        con = sqlite3.connect(db_path)
        con.row_factory = sqlite3.Row
        cur = con.execute(q)
        if cur.description:
            cols = [d[0] for d in cur.description]
            rows = [dict(r) for r in cur.fetchmany(200)]
        else:
            cols, rows = [], []
        if allow_write:
            con.commit()
        affected = cur.rowcount
        con.close()
        return {"ok": True, "columns": cols, "rows": rows, "row_count": len(rows), "affected": affected}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}


_TODO_STATUSES = {"pending", "in_progress", "completed"}


def todo_write(todos: list) -> dict:
    """Mantiene la lista de tareas (TODOs) visible para el usuario en tareas largas,
    igual que Claude Code. Llama esto al planificar y cada vez que cambie el estado
    de un paso. Cada todo: {content, status} con status pending|in_progress|completed.
    Regla: solo UN paso in_progress a la vez; marca completed en cuanto termines uno."""
    if not isinstance(todos, list) or not todos:
        return {"ok": False, "error": "todos debe ser una lista no vacía de {content,status}"}
    clean = []
    for i, t in enumerate(todos):
        if not isinstance(t, dict) or not t.get("content"):
            return {"ok": False, "error": f"todo #{i+1} inválido: falta content"}
        status = str(t.get("status", "pending")).strip().lower()
        if status not in _TODO_STATUSES:
            status = "pending"
        clean.append({"content": str(t["content"]).strip()[:200], "status": status})
    done = sum(1 for t in clean if t["status"] == "completed")
    return {"ok": True, "todos": clean, "summary": f"{done}/{len(clean)} completadas"}


def multi_edit(path: str, edits: list) -> dict:
    """Aplica VARIAS ediciones a un archivo en una sola llamada, de forma ATÓMICA:
    si alguna falla, NO se escribe nada (todo o nada). Cada edit es un dict
    {old_string, new_string, replace_all?}. Se aplican en orden sobre el resultado
    de las anteriores. Devuelve un único diff combinado. Ideal para refactors."""
    if not os.path.isfile(path):
        return {"ok": False, "error": "no existe el archivo: " + str(path)}
    if not isinstance(edits, list) or not edits:
        return {"ok": False, "error": "edits debe ser una lista no vacía de {old_string,new_string}"}
    try:
        original, eol = _read_with_eol(path)
    except Exception as e:
        return {"ok": False, "error": "lectura: " + str(e)}

    content = original
    total = 0
    modes = []
    for i, ed in enumerate(edits):
        if not isinstance(ed, dict) or "old_string" not in ed or "new_string" not in ed:
            return {"ok": False, "error": f"edit #{i+1} inválido: faltan old_string/new_string"}
        try:
            content, count, mode = _apply_replacement(
                content, ed["old_string"], ed["new_string"], bool(ed.get("replace_all", False)))
        except ValueError as e:
            return {"ok": False, "error": f"edit #{i+1} falló: {e}. No se escribió nada (atómico)."}
        total += count
        modes.append(mode)

    diff = _make_unified_diff(original, content, path)
    try:
        size = _write_with_eol(path, content, eol)
    except Exception as e:
        return {"ok": False, "error": "escritura: " + str(e)}
    result = {"ok": True, "path": path, "edits_applied": len(edits),
              "replacements": total, "new_size": size, "match_modes": modes, "diff": diff}
    return _py_syntax_result(path, content, result)
